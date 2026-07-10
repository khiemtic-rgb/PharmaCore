#!/usr/bin/env python3
"""Xuat Markdown -> Word (.docx) cho ho so CQD. Yeu cau: pip install python-docx markdown"""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Inches, Pt
except ImportError:
    print("Cai dat: pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "word"
SCREENSHOTS = ROOT / "assets" / "screenshots"

DOCS = [
    ("01-software-functional-specification-v1.md", "NVX-CQD-SFS-01.docx"),
    ("02-api-integration-specification-v1.md", "NVX-CQD-API-01.docx"),
    ("03-system-architecture-document-v1.md", "NVX-CQD-ARCH-01.docx"),
    ("README.md", "NVX-CQD-00-Muc-luc-ho-so.docx"),
    ("phu-luc-a-qd540-field-map-v1.md", "NVX-CQD-PL-A-QD540-Field-Map.docx"),
]


def set_default_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int) -> None:
    clean = text.strip().strip("#").strip()
    doc.add_heading(clean, level=min(level, 4))


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line.strip()))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell_text = row[j] if j < len(row) else ""
            table.rows[i].cells[j].text = cell_text


def resolve_image(path_str: str, md_dir: Path) -> Path | None:
    path_str = path_str.strip()
    if path_str.startswith("http://") or path_str.startswith("https://"):
        return None
    candidates = [
        (md_dir / path_str).resolve(),
        (ROOT / path_str).resolve(),
        (SCREENSHOTS / Path(path_str).name).resolve(),
    ]
    for c in candidates:
        if c.exists() and c.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}:
            return c
    return None


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    set_default_style(doc)

    lines = text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        # fenced code
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # table
        if line.strip().startswith("|"):
            if is_table_separator(line):
                i += 1
                continue
            table_buf.append(parse_table_row(line))
            i += 1
            # peek next
            if i >= len(lines) or not lines[i].strip().startswith("|"):
                add_table(doc, table_buf)
                table_buf = []
            continue

        if table_buf:
            add_table(doc, table_buf)
            table_buf = []

        # image ![alt](path)
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
        if img_match:
            alt, src = img_match.group(1), img_match.group(2)
            img_path = resolve_image(src, md_path.parent)
            if img_path:
                doc.add_picture(str(img_path), width=Inches(6.0))
                cap = doc.add_paragraph(alt or img_path.name)
                cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            else:
                doc.add_paragraph(f"[Anh: {alt} — {src}]")
            i += 1
            continue

        # headings
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            add_heading(doc, line[level:], level)
            i += 1
            continue

        # horizontal rule
        if re.match(r"^-{3,}$", line.strip()):
            doc.add_paragraph("—" * 40)
            i += 1
            continue

        # empty
        if not line.strip():
            i += 1
            continue

        # bullet
        if re.match(r"^[-*]\s+", line.strip()):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
            i += 1
            continue

        # numbered list
        if re.match(r"^\d+\.\s+", line.strip()):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line.strip()), style="List Number")
            i += 1
            continue

        # blockquote
        if line.strip().startswith(">"):
            doc.add_paragraph(line.strip().lstrip(">").strip())
            i += 1
            continue

        # normal paragraph (strip simple md bold/inline code)
        plain = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
        plain = re.sub(r"`([^`]+)`", r"\1", plain)
        plain = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", plain)
        doc.add_paragraph(plain)
        i += 1

    if table_buf:
        add_table(doc, table_buf)
    if code_buf:
        add_code_block(doc, code_buf)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"  OK {docx_path.name}")


def patch_sfs_with_images() -> None:
    """Chen link anh vao SFS neu chua co."""
    sfs = ROOT / "01-software-functional-specification-v1.md"
    content = sfs.read_text(encoding="utf-8")
    shots = [
        ("01-login.png", "Man hinh dang nhap"),
        ("02-dashboard.png", "Dashboard tong quan"),
        ("03-pos.png", "POS ban hang"),
        ("04-drug-master.png", "Danh muc thuoc"),
        ("05-inventory.png", "Quan ly kho FEFO"),
        ("06-grn.png", "Nhap thuoc GRN"),
        ("07-sale.png", "Don ban thuoc"),
        ("08-report-nxt.png", "Bao cao nhap xuat ton"),
        ("09-report-revenue.png", "Bao cao doanh thu"),
        ("10-drug-connectivity.png", "Lien thong QD540"),
        ("11-api-config.png", "Cau hinh API CSDL Duoc QG"),
        ("12-sync-log.png", "Nhat ky dong bo"),
    ]
    block = "\n".join(
        f"![{label}](./assets/screenshots/{fname})" for fname, label in shots
    )
    marker = "## 10. Hình ảnh giao diện minh họa"
    if marker in content and "![Man hinh dang nhap]" not in content and "![Màn hình đăng nhập]" not in content:
        # Replace table-only section with images + table
        pattern = r"(## 10\. Hình ảnh giao diện minh họa\n\n)(Tài liệu đính kèm.*?\n\n\| STT \|.*?\| `12-sync-log\.png` \|\n\n---)"
        replacement = rf"\1\n{block}\n\n---"
        new_content = re.sub(pattern, replacement, content, flags=re.DOTALL)
        if new_content != content:
            sfs.write_text(new_content, encoding="utf-8")


def main() -> None:
    patch_sfs_with_images()
    print(f"Xuat Word -> {OUT_DIR}")
    for md_name, docx_name in DOCS:
        md_path = ROOT / md_name
        if not md_path.exists():
            print(f"  SKIP missing {md_name}")
            continue
        convert_md_to_docx(md_path, OUT_DIR / docx_name)
    print("Hoan tat.")


if __name__ == "__main__":
    main()
